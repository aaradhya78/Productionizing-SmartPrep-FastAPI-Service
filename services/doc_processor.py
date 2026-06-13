import io
import re
import logging
from typing import List, Dict, Any

logger = logging.getLogger("ai_service")

# Try importing docx and pytesseract
try:
    import docx
except ImportError:
    docx = None
    logger.warning("python-docx is not installed or import failed.")

try:
    import pytesseract
    from PIL import Image
    PYTESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None
    PYTESSERACT_AVAILABLE = False
    logger.warning("pytesseract or PIL is not installed or import failed. OCR support disabled.")

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from file bytes based on file format (PDF, DOCX, TXT) with OCR fallback for scanned PDFs."""
    ext = filename.split(".")[-1].lower()
    
    if ext == "pdf":
        return extract_pdf(file_bytes)
    elif ext == "docx":
        return extract_docx(file_bytes)
    elif ext in ["txt", "md", "csv"]:
        return extract_txt(file_bytes)
    else:
        # Default fallback
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""

def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF, applying OCR to scanned pages if necessary."""
    import fitz  # PyMuPDF
    
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    extracted_text = ""
    is_scanned = True
    
    # First pass: try extracting standard embedded text
    pages_text = []
    for page in doc:
        text = page.get_text()
        pages_text.append(text)
    
    total_len = sum(len(t.strip()) for t in pages_text)
    
    # If the text is substantial, we assume it's a digital PDF
    if total_len > 200:
        is_scanned = False
        extracted_text = "\n".join(pages_text)
    
    # If scanned and we have OCR capabilities, run OCR
    if is_scanned:
        logger.info("PDF appears to be scanned or image-based. Attempting OCR...")
        if PYTESSERACT_AVAILABLE:
            ocr_pages = []
            for i, page in enumerate(doc):
                try:
                    # Render page as image (pixmap)
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Run OCR
                    page_text = pytesseract.image_to_string(img)
                    ocr_pages.append(page_text)
                except Exception as e:
                    logger.error(f"OCR failed for page {i+1}: {e}")
                    # Fallback to standard text on failure
                    ocr_pages.append(page.get_text())
            extracted_text = "\n".join(ocr_pages)
        else:
            logger.warning("OCR requested but pytesseract/Pillow is not available. Returning raw extracted text.")
            extracted_text = "\n".join(pages_text)
            
    return extracted_text

def extract_docx(file_bytes: bytes) -> str:
    """Extract text from Microsoft Word (DOCX) files."""
    if not docx:
        logger.error("DOCX library missing, cannot extract docx.")
        return "Error: docx library not installed."
        
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n\n".join(full_text)

def extract_txt(file_bytes: bytes) -> str:
    """Extract text from plain text files."""
    for encoding in ["utf-8", "latin-1", "utf-16"]:
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")

def clean_text(text: str) -> str:
    """Clean extracted text: remove headers, footers, page numbers, duplicate lines, and normalize spacing."""
    if not text:
        return ""
        
    lines = text.split("\n")
    cleaned_lines = []
    
    # 1. Regex patterns for page numbers, headers, and footers
    page_num_patterns = [
        re.compile(r"^\s*page\s*\d+\s*(of\s*\d+)?\s*$", re.IGNORECASE),
        re.compile(r"^\s*pg\.?\s*\d+\s*$", re.IGNORECASE),
        re.compile(r"^\s*\d+\s*of\s*\d+\s*$", re.IGNORECASE),
        re.compile(r"^\s*-\s*\d+\s*-\s*$", re.IGNORECASE),
        re.compile(r"^\s*\d+\s*$")  # solitary digits on a line
    ]
    
    confidential_patterns = [
        re.compile(r"confidential", re.IGNORECASE),
        re.compile(r"all rights reserved", re.IGNORECASE),
        re.compile(r"copyright \u00a9|copyright ©", re.IGNORECASE),
        re.compile(r"^draft$", re.IGNORECASE)
    ]

    # Heuristic to remove repeating headers/footers: 
    # we collect lines and count frequencies of short lines that appear at the beginning or end of pages.
    # For simplicity, we filter out line-by-line using regex first.
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            cleaned_lines.append("")
            continue
            
        # Check if page number
        is_page_num = any(pat.match(line_strip) for pat in page_num_patterns)
        if is_page_num:
            continue
            
        # Check if copyright/confidential footer
        is_confidential = any(pat.search(line_strip) for pat in confidential_patterns)
        if is_confidential:
            continue
            
        cleaned_lines.append(line)
        
    cleaned_text = "\n".join(cleaned_lines)
    
    # 2. Normalize whitespace and newlines
    cleaned_text = re.sub(r"[ \t]+", " ", cleaned_text)  # collapse spaces/tabs
    cleaned_text = re.sub(r"\n{3,}", "\n\n", cleaned_text)  # limit consecutive newlines to 2
    
    return cleaned_text.strip()

def chunk_text(text: str, chunk_size: int = 1500, overlap: int = 200) -> List[str]:
    """Split clean text into semantic chunks (~chunk_size characters) with overlap."""
    if not text:
        return []
        
    if len(text) <= chunk_size:
        return [text]
        
    # Split text into paragraphs
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        # If a single paragraph is larger than chunk_size, split it into sentences
        if len(para) > chunk_size:
            # First commit what we have so far
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_length = 0
                
            # Split paragraph into sentences
            sentences = re.split(r"(?<=[.!?])\s+", para)
            current_sentence_chunk = []
            current_sentence_len = 0
            
            for sentence in sentences:
                if len(sentence) > chunk_size:
                    # If sentence itself is huge, force split it by character chunks
                    if current_sentence_chunk:
                        chunks.append(" ".join(current_sentence_chunk))
                        current_sentence_chunk = []
                        current_sentence_len = 0
                    
                    # Force chunking
                    for i in range(0, len(sentence), chunk_size - overlap):
                        chunks.append(sentence[i:i + chunk_size])
                elif current_sentence_len + len(sentence) > chunk_size:
                    chunks.append(" ".join(current_sentence_chunk))
                    # Retain overlap sentences
                    overlap_len = 0
                    overlap_chunk = []
                    # Backtrack for overlap
                    for s in reversed(current_sentence_chunk):
                        if overlap_len + len(s) < overlap:
                            overlap_chunk.insert(0, s)
                            overlap_len += len(s) + 1
                        else:
                            break
                    current_sentence_chunk = overlap_chunk + [sentence]
                    current_sentence_len = overlap_len + len(sentence) + 1
                else:
                    current_sentence_chunk.append(sentence)
                    current_sentence_len += len(sentence) + 1
                    
            if current_sentence_chunk:
                current_chunk = current_sentence_chunk
                current_length = current_sentence_len
        elif current_length + len(para) > chunk_size:
            chunks.append("\n\n".join(current_chunk))
            
            # Start new chunk with overlap paragraphs if possible
            overlap_len = 0
            overlap_chunk = []
            for p in reversed(current_chunk):
                if overlap_len + len(p) < overlap:
                    overlap_chunk.insert(0, p)
                    overlap_len += len(p) + 2
                else:
                    break
            current_chunk = overlap_chunk + [para]
            current_length = overlap_len + len(para) + 2
        else:
            current_chunk.append(para)
            current_length += len(para) + 2
            
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))
        
    return chunks
